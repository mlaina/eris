"""Build Eris-v0.docx from the markdown sources.

Structure:
  - Cover page (title)
  - Map (PNG)
  - 24 chapters (acto-I → acto-V)
  - Glosario
  - Lugares
  - Personajes
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
CAPS_DIR = ROOT / "estructura" / "capitulos"
CANON_DIR = ROOT / "canon"
MAP_PNG = ROOT / "build" / "mapa.png"
OUT = ROOT / "build" / "Eris-v0.docx"

CHAPTERS = [
    ("acto-I", [
        "01-prologo.md",
        "02-elis-lirae.md",
        "03-plaza-ocre.md",
        "04-libro-pequeno.md",
    ]),
    ("acto-II", [
        "05-audiencia-larien.md",
        "06-nurn-dia.md",
        "07-cena-elerian.md",
        "08-atentado.md",
        "09-mision-135.md",
        "10-pesadillas.md",
    ]),
    ("acto-III", [
        "11-corte-hueca.md",
        "12-reunion-manos.md",
        "13-bajo-palacio.md",
        "14-lirae-pared.md",
        "15-aislamiento.md",
    ]),
    ("acto-IV", [
        "16-cena-confesiones.md",
        "17-huida-haris.md",
        "18-viaje-ygoran.md",
        "19-ygoran-rebelde.md",
        "20-descubrimiento-lirae.md",
    ]),
    ("acto-V", [
        "21-tejado-ygoran.md",
        "22-asalto-ygoran.md",
        "23-el-muelle.md",
        "24-epilogo-nyree.md",
    ]),
]

ACT_TITLES = {
    "acto-I": "Acto I — Gloria",
    "acto-II": "Acto II — Grietas",
    "acto-III": "Acto III — Revelación",
    "acto-IV": "Acto IV — Fuga",
    "acto-V": "Acto V — Tragedia",
}

# Personajes: orden manual + filtrado de índice/grupos.
PERSONAJES_ORDER = [
    "eris.md",
    "haris.md",
    "muoro.md",
    "theim.md",
    "faree.md",
    "lyris.md",
    "lirae.md",
    "elerian.md",
    "nyree.md",
    "velerian.md",
    "maelor.md",
]

# Lugares: ordenados por relevancia narrativa
LUGARES_ORDER = [
    "coria-elis.md",
    "nurn.md",
    "itsu.md",
    "ygoran.md",
    "kaurun.md",
    "liga-kovax.md",
    "indra.md",
    "sandrie.md",
    "aret.md",
    "fireon.md",
    "deengar.md",
    "asnare.md",
    "geoda.md",
    "sacrio.md",
    "saebal.md",
    "theris.md",
    "trei.md",
    "valdara.md",
    "zonas-extasiadas.md",
    "anira-geografia.md",
]


def strip_obsidian_links(text: str) -> str:
    """Resolve [[archivo|título]] → título, [[archivo]] → archivo (last segment)."""
    def repl(m):
        inner = m.group(1)
        if "|" in inner:
            return inner.split("|", 1)[1]
        # Use the last path segment as visible text
        return inner.rsplit("/", 1)[-1].replace("-", " ")
    return re.sub(r"\[\[([^\]]+)\]\]", repl, text)


def strip_canon_block(content: str) -> str:
    """Cortar el bloque '## Canon relacionado' (metadata) y posteriores."""
    # Cortar a partir de "## Canon relacionado" o "## Pendientes" o "## Sesión"
    for marker in ["## Canon relacionado", "## Pendientes", "## Sesión", "## Notas"]:
        idx = content.find(marker)
        if idx != -1:
            content = content[:idx]
    return content.rstrip()


def add_horizontal_rule(paragraph):
    """Insertar separador horizontal en un párrafo."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline_runs(paragraph, text: str):
    """Añadir runs procesando **bold** y *italic* básicos."""
    text = strip_obsidian_links(text)
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_markdown_body(doc, content: str, base_heading_level: int = 2):
    """Convertir markdown a párrafos Word.

    base_heading_level: nivel a aplicar al primer `#`. Los headings dentro del
    cuerpo (raros) se ajustan a partir de ahí.
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip the document's H1 (it ya se inserta como título de capítulo aparte)
        if line.startswith("# "):
            i += 1
            continue

        if not line:
            i += 1
            continue

        if line == "---":
            # Separador
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("⁂")
            run.font.size = Pt(14)
            i += 1
            continue

        # Headings
        h_match = re.match(r"^(#{2,6})\s+(.*)$", line)
        if h_match:
            level = len(h_match.group(1))
            text = strip_obsidian_links(h_match.group(2))
            doc.add_heading(text, level=min(level + base_heading_level - 2, 9))
            i += 1
            continue

        # Bullet list
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
            i += 1
            continue

        # Numbered list (1. 2. ...)
        num_match = re.match(r"^\d+\.\s+(.*)$", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, num_match.group(1))
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Garamond")
    rFonts.set(qn("w:hAnsi"), "Garamond")


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc):
    # Espacio vertical
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ERIS")
    run.font.size = Pt(48)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tragedia política y militar de Anira")
    run.font.size = Pt(14)
    run.italic = True

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Marcos Laina")
    run.font.size = Pt(16)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión v0 — 2026-05-26")
    run.font.size = Pt(10)
    run.italic = True
    page_break(doc)


def add_map(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mapa de Anira")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph()

    if MAP_PNG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Ajustar al ancho de página manteniendo proporción.
        run.add_picture(str(MAP_PNG), width=Inches(6.5))
    page_break(doc)


def add_chapter(doc, act_dir, filename):
    path = CAPS_DIR / act_dir / filename
    raw = path.read_text(encoding="utf-8")
    raw = strip_canon_block(raw)

    # Sacar el título (primer #)
    title_match = re.match(r"^#\s+(.*)$", raw, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else filename

    doc.add_heading(title, level=1)
    add_markdown_body(doc, raw, base_heading_level=2)
    page_break(doc)


def add_canon_section(doc, section_title, file_list, base_dir):
    doc.add_heading(section_title, level=1)
    for filename in file_list:
        path = base_dir / filename
        if not path.exists():
            continue
        raw = path.read_text(encoding="utf-8")
        raw = strip_canon_block(raw)
        # Título de la ficha
        title_match = re.match(r"^#\s+(.*)$", raw, re.MULTILINE)
        title = title_match.group(1).strip() if title_match else filename
        doc.add_heading(title, level=2)
        add_markdown_body(doc, raw, base_heading_level=3)
        doc.add_paragraph()
    page_break(doc)


def add_glosario(doc):
    path = CANON_DIR / "glosario.md"
    if not path.exists():
        return
    raw = path.read_text(encoding="utf-8")
    raw = strip_canon_block(raw)
    title_match = re.match(r"^#\s+(.*)$", raw, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else "Glosario"
    doc.add_heading(title, level=1)
    add_markdown_body(doc, raw, base_heading_level=2)
    page_break(doc)


def main():
    doc = Document()
    set_default_font(doc)

    add_cover(doc)
    add_map(doc)

    for act_dir, files in CHAPTERS:
        # Heading de acto al inicio
        doc.add_heading(ACT_TITLES[act_dir], level=1)
        page_break(doc)
        for filename in files:
            add_chapter(doc, act_dir, filename)

    add_glosario(doc)
    add_canon_section(doc, "Lugares", LUGARES_ORDER, CANON_DIR / "lugares")
    add_canon_section(doc, "Personajes", PERSONAJES_ORDER, CANON_DIR / "personajes")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
